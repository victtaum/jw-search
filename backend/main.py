import uvicorn
from typing import Optional, List, Dict, Any
from fastapi import FastAPI, Query, Header, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import os

from pydantic import BaseModel
from scraper import perform_ai_grounded_search, get_clean_document, set_api_key, get_api_status, ApiKeyRequiredException
from rag_engine import perform_custom_rag_search

class KeyConfigRequest(BaseModel):
    api_key: str

app = FastAPI(
    title="JW Search API",
    description="Backend de consulta de informações do jw.org e wol.jw.org com suporte a Inteligência Artificial",
    version="1.4.0"
)

# Configure CORS so both local web frontend and Android app can access the API
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

import re
from io import BytesIO
import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi.responses import Response

class ChatMessage(BaseModel):
    role: str # "user" or "assistant"
    content: str

class ChatRequest(BaseModel):
    query: str
    history: List[ChatMessage] = []
    provider: Optional[str] = "hy3"
    model: Optional[str] = None
    base_url: Optional[str] = None
    include_external: bool = False
    lang: str = "pt"

ChatRequest.model_rebuild()

class ExportDocxRequest(BaseModel):
    title: str = "Estudo Teocrático - JW Search"
    content: str


def create_theocratic_docx(title: str, markdown_content: str) -> BytesIO:
    doc = docx.Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Document Header
    header = doc.add_heading(title, level=0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if header.runs:
        header.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        header.runs[0].font.name = "Calibri"

    p_sub = doc.add_paragraph("Documento de Estudo Bíblico gerado via JW Search (Fontes: wol.jw.org e jw.org)")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if p_sub.runs:
        p_sub.runs[0].font.italic = True
        p_sub.runs[0].font.size = Pt(9.5)
        p_sub.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph() # Spacer

    # Simple Markdown parser for Docx
    lines = markdown_content.split("\n")
    in_table = False
    table_rows = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if in_table and table_rows:
                # Render table
                _render_docx_table(doc, table_rows)
                in_table = False
                table_rows = []
            continue

        # Check Table row
        if stripped.startswith("|") and stripped.endswith("|"):
            if "---" in stripped:
                continue # Header divider line
            in_table = True
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            table_rows.append(cells)
            continue
        elif in_table and table_rows:
            _render_docx_table(doc, table_rows)
            in_table = False
            table_rows = []

        # Headings
        if stripped.startswith("### "):
            h = doc.add_heading(stripped[4:], level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(30, 58, 138)
        elif stripped.startswith("## "):
            h = doc.add_heading(stripped[3:], level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(15, 23, 42)
        elif stripped.startswith("# "):
            h = doc.add_heading(stripped[2:], level=0)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(15, 23, 42)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            _add_markdown_runs(p, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Number')
            _add_markdown_runs(p, re.sub(r'^\d+\.\s', '', stripped))
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            _add_markdown_runs(p, stripped.lstrip("> "))
            if p.runs:
                p.runs[0].font.italic = True
                p.runs[0].font.color.rgb = RGBColor(71, 85, 105)
        else:
            p = doc.add_paragraph()
            _add_markdown_runs(p, stripped)

    if in_table and table_rows:
        _render_docx_table(doc, table_rows)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def _add_markdown_runs(paragraph, text):
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("[") and "](" in part and part.endswith(")"):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                label, url = m.groups()
                run = paragraph.add_run(f"{label} ({url})")
                run.font.color.rgb = RGBColor(37, 99, 235)
                run.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def _render_docx_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Light Shading Accent 1' if 'Light Shading Accent 1' in [s.name for s in doc.styles] else 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < col_count:
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
    doc.add_paragraph() # Spacing after table


@app.post("/api/chat")
def api_chat(
    req: ChatRequest,
    x_gemini_api_key: Optional[str] = Header(None, alias="X-Gemini-Api-Key"),
    x_deepseek_api_key: Optional[str] = Header(None, alias="X-Deepseek-Api-Key"),
    x_hy3_api_key: Optional[str] = Header(None, alias="X-Hy3-Api-Key"),
    x_api_key: Optional[str] = Header(None, alias="X-Api-Key"),
):
    return handle_theocratic_search(
        q=req.query,
        external=req.include_external,
        lang=req.lang,
        provider=req.provider,
        model=req.model,
        base_url=req.base_url,
        history=[{"role": m.role, "content": m.content} for m in req.history],
        x_gemini_api_key=x_gemini_api_key,
        x_deepseek_api_key=x_deepseek_api_key,
        x_hy3_api_key=x_hy3_api_key,
        x_api_key=x_api_key
    )


@app.post("/api/export/docx")
def api_export_docx(data: ExportDocxRequest):
    if not data.content.strip():
        raise HTTPException(status_code=400, detail="Conteúdo vazio para exportação.")
    try:
        docx_io = create_theocratic_docx(data.title, data.content)
        return Response(
            content=docx_io.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="estudo_teocratico_{data.title[:20].replace(" ", "_")}.docx"'
            }
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao gerar DOCX: {e}")


def handle_theocratic_search(
    q: str,
    external: bool = False,
    lang: str = "pt",
    provider: str = "hy3",
    model: Optional[str] = None,
    base_url: Optional[str] = None,
    history: Optional[List[Dict[str, str]]] = None,
    x_gemini_api_key: Optional[str] = None,
    x_deepseek_api_key: Optional[str] = None,
    x_hy3_api_key: Optional[str] = None,
    x_api_key: Optional[str] = None,
    api_key: Optional[str] = None
):
    if not q.strip():
        return {"ai_response": "", "results": []}
    
    prov = (provider or "gemini").lower()
    
    # 1. Custom Theocratic RAG for DeepSeek
    if prov == "deepseek":
        active_key = x_deepseek_api_key or x_api_key or api_key or os.environ.get("DEEPSEEK_API_KEY")
        if not active_key:
            # Fallback to Gemini or Hy3 if available
            gemini_key = x_gemini_api_key or os.environ.get("GEMINI_API_KEY")
            if gemini_key:
                return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=gemini_key, history=history)
            hy3_key = x_hy3_api_key or os.environ.get("HY3_API_KEY") or os.environ.get("OPENAI_API_KEY")
            if hy3_key:
                return perform_custom_rag_search(query=q.strip(), provider="hy3", api_key=hy3_key, include_external=external, lang=lang, history=history)
            raise HTTPException(
                status_code=401,
                detail="Chave da API do DeepSeek não informada. Adicione sua chave nas configurações para pesquisar."
            )
        try:
            res = perform_custom_rag_search(
                query=q.strip(),
                provider="deepseek",
                api_key=active_key,
                base_url=base_url,
                model=model,
                include_external=external,
                lang=lang,
                history=history
            )
            if res.get("ai_response") and len(res["ai_response"].strip()) > 20:
                return res
            # If DeepSeek produced empty output, fallback to Gemini
            gemini_key = x_gemini_api_key or os.environ.get("GEMINI_API_KEY")
            if gemini_key:
                return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=gemini_key, history=history)
            return res
        except Exception as e:
            try:
                gemini_key = x_gemini_api_key or os.environ.get("GEMINI_API_KEY")
                return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=gemini_key, history=history)
            except Exception:
                raise HTTPException(status_code=500, detail=str(e))

    # 2. Hy3 / Tencent / OpenRouter (with automatic fallback to Gemini)
    elif prov in ["hy3", "tencent", "hunyuan", "openai"]:
        active_key = x_hy3_api_key or x_api_key or api_key or os.environ.get("HY3_API_KEY") or os.environ.get("OPENAI_API_KEY")
        gemini_key = x_gemini_api_key or os.environ.get("GEMINI_API_KEY")
        has_gemini = bool(gemini_key)
        
        if not active_key:
            # If Hy3 has no key, try Gemini fallback immediately
            if has_gemini:
                try:
                    return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=gemini_key, history=history)
                except Exception as g_err:
                    pass
            raise HTTPException(
                status_code=401,
                detail="Chave da API (Hy3/OpenRouter) não informada. Adicione sua chave nas configurações para pesquisar."
            )
        try:
            res = perform_custom_rag_search(
                query=q.strip(),
                provider=prov,
                api_key=active_key,
                base_url=base_url,
                model=model,
                include_external=external,
                lang=lang,
                history=history
            )
            if res.get("ai_response") and len(res["ai_response"].strip()) > 20:
                return res
            # If Hy3 returned empty text, transparently fallback to Gemini
            if has_gemini:
                print("Hy3 returned empty response. Falling back to Google Gemini Grounding...")
                return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=gemini_key, history=history)
            return res
        except Exception as e:
            if has_gemini:
                print(f"Hy3 failed ({e}). Attempting automatic fallback to Google Gemini...")
                try:
                    return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=gemini_key, history=history)
                except Exception as g_err:
                    pass
            raise HTTPException(
                status_code=429 if "429" in str(e) else 500,
                detail="As cotas padrão do Hy3 e do Gemini falharam. Adicione sua própria chave gratuita no painel de configurações para continuar pesquisando sem limites."
            )

    # 3. Google Gemini with Grounding (with auto-fallback to Hy3)
    else:
        client_key = x_gemini_api_key or x_api_key or api_key
        hy3_key = x_hy3_api_key or os.environ.get("HY3_API_KEY") or os.environ.get("OPENAI_API_KEY")
        try:
            return perform_ai_grounded_search(q.strip(), include_external=external, lang=lang, custom_api_key=client_key, history=history)
        except ApiKeyRequiredException:
            if hy3_key:
                print("Gemini key required. Falling back to Hy3 RAG...")
                return perform_custom_rag_search(query=q.strip(), provider="hy3", api_key=hy3_key, include_external=external, lang=lang, history=history)
            raise HTTPException(status_code=401, detail="Chave do Google Gemini não informada. Adicione sua chave nas configurações.")
        except Exception as e:
            err_str = str(e)
            if "429" in err_str or "RESOURCE_EXHAUSTED" in err_str or "quota" in err_str.lower() or "500" in err_str:
                if hy3_key:
                    print("Gemini failed/quota exhausted. Automatically falling back to Hy3 RAG engine...")
                    return perform_custom_rag_search(
                        query=q.strip(),
                        provider="hy3",
                        api_key=hy3_key,
                        include_external=external,
                        lang=lang,
                        history=history
                    )
                raise HTTPException(
                    status_code=429,
                    detail="Cota da API do Gemini esgotada. Configure sua chave gratuita do Hy3 ou DeepSeek nas configurações."
                )
            raise HTTPException(status_code=500, detail=err_str)


@app.get("/api/search")
def api_search(
    q: str = Query(..., description="Termo de pesquisa"),
    external: bool = Query(False, description="Incluir fontes externas da internet"),
    lang: str = Query("pt", description="Código de idioma (ex: pt, en, es)"),
    provider: str = Query("hy3", description="Provedor de IA: hy3, gemini, deepseek, openai"),
    model: Optional[str] = Query(None, description="Modelo específico"),
    base_url: Optional[str] = Query(None, description="Endpoint base customizado da API"),
    x_gemini_api_key: Optional[str] = Header(None, alias="X-Gemini-Api-Key"),
    x_deepseek_api_key: Optional[str] = Header(None, alias="X-Deepseek-Api-Key"),
    x_hy3_api_key: Optional[str] = Header(None, alias="X-Hy3-Api-Key"),
    x_api_key: Optional[str] = Header(None, alias="X-Api-Key"),
    api_key: Optional[str] = Query(None, description="Chave API opcional do cliente")
):
    return handle_theocratic_search(
        q=q,
        external=external,
        lang=lang,
        provider=provider,
        model=model,
        base_url=base_url,
        history=None,
        x_gemini_api_key=x_gemini_api_key,
        x_deepseek_api_key=x_deepseek_api_key,
        x_hy3_api_key=x_hy3_api_key,
        x_api_key=x_api_key,
        api_key=api_key
    )



@app.get("/api/read")
def api_read(
    url: str = Query(..., description="URL absoluta do documento wol.jw.org")
):
    if not url.strip():
        raise HTTPException(status_code=400, detail="URL inválida")
    try:
        content = get_clean_document(url.strip())
        if not content:
            raise HTTPException(status_code=404, detail="Documento não encontrado ou erro ao acessar")
        return {"content": content}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/config")
def api_get_config():
    return get_api_status()

@app.get("/api/diagnostics")
def api_diagnostics():
    import time
    report = {
        "status": "online",
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S UTC", time.gmtime()),
        "providers": {}
    }
    
    # 1. Test WOL scraper speed
    t0 = time.time()
    try:
        wol_res = search_wol_direct("amor leal", lang="pt", max_results=3)
        wol_time = round((time.time() - t0) * 1000, 1)
        report["providers"]["wol_library"] = {
            "status": "ok",
            "latency_ms": wol_time,
            "sample_results": len(wol_res)
        }
    except Exception as e:
        report["providers"]["wol_library"] = {"status": "error", "error": str(e)}

    # 2. Check Gemini config
    gemini_key = os.environ.get("GEMINI_API_KEY")
    report["providers"]["gemini"] = {
        "configured": bool(gemini_key),
        "status": "ready" if gemini_key else "not_configured"
    }

    # 3. Check Hy3 config
    hy3_key = os.environ.get("HY3_API_KEY") or os.environ.get("OPENAI_API_KEY")
    report["providers"]["hy3_openrouter"] = {
        "configured": bool(hy3_key),
        "status": "ready" if hy3_key else "not_configured"
    }

    # 4. Check DeepSeek config
    deepseek_key = os.environ.get("DEEPSEEK_API_KEY")
    report["providers"]["deepseek"] = {
        "configured": bool(deepseek_key),
        "status": "ready" if deepseek_key else "not_configured"
    }

    return report

@app.post("/api/config")
def api_set_config(data: KeyConfigRequest):
    success, msg = set_api_key(data.api_key)
    if not success:
        raise HTTPException(status_code=400, detail=msg)
    return {"message": msg, "has_key": True}

# Locate and mount static files to serve the web frontend
web_candidates = [
    os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "web")),
    os.path.abspath(os.path.join(os.path.dirname(__file__), "web")),
    os.path.abspath("web")
]
web_dir = next((p for p in web_candidates if os.path.exists(p)), None)

if web_dir:
    app.mount("/", StaticFiles(directory=web_dir, html=True), name="web")
else:
    @app.get("/")
    def read_root():
        return {
            "message": "JW Search API está rodando. O diretório do frontend web '/web' não foi encontrado para ser servido na raiz.",
            "api_search": "/api/search?q=termo&external=false",
            "api_read": "/api/read?url=https://wol.jw.org/pt/wol/d/r5/lp-t/1200002781"
        }

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run("main:app", host="0.0.0.0", port=port)

